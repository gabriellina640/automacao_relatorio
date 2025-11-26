import sys
import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from datetime import datetime

# ==============================================================================
# 1. FUNÇÕES AUXILIARES DE ESTILO
# ==============================================================================

def desativar_corretor(run):
    """Insere a tag XML w:noProof para não sublinhar de vermelho."""
    rPr = run._element.get_or_add_rPr()
    noProof = OxmlElement("w:noProof")
    rPr.append(noProof)

def formatar_paragrafo_hibrido(paragrafo, texto_resolucao, texto_assunto):
    """
    Cria um parágrafo com formatação mista:
    - Resolução: Arial 12, Negrito, Sublinhado.
    - Assunto: Arial 12, Normal.
    """
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragrafo.paragraph_format.space_after = Pt(12) 
    paragrafo.paragraph_format.line_spacing = 1.15  
    
    # 1. Parte da Resolução (Negrito + Sublinhado)
    run_res = paragrafo.add_run(texto_resolucao)
    run_res.font.name = 'Arial'
    run_res.font.size = Pt(12)
    run_res.font.bold = True
    run_res.font.underline = True
    desativar_corretor(run_res)
    
    # 2. Separador e Assunto (Normal)
    run_assunto = paragrafo.add_run(f" - {texto_assunto}")
    run_assunto.font.name = 'Arial'
    run_assunto.font.size = Pt(12)
    run_assunto.font.bold = False
    run_assunto.font.underline = False
    desativar_corretor(run_assunto)

def formatar_celula_tabela(cell, texto, negrito=False):
    """(Usado apenas na Pauta - Botão 1 - Mantém Times 8)"""
    cell.text = str(texto) if texto else ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0 
        
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(8)
            run.font.bold = negrito
            desativar_corretor(run)

# ==============================================================================
# 2. LÓGICA DO BOTÃO 1: PAUTA (Mantida Times 8 - Tabela)
# ==============================================================================

def limpar_nome_sessao(texto_completo):
    match_num = re.search(r"(\d+[ªº°])", texto_completo)
    numero = match_num.group(1) if match_num else ""
    tipo = "Virtual" if "Virtual" in texto_completo else ""
    resultado = f"{numero} {tipo}".strip()
    return resultado if resultado else "Sessão"

def encontrar_sessao_formatada(doc):
    for i in range(min(15, len(doc.paragraphs))):
        texto = doc.paragraphs[i].text.strip()
        if "Sessão" in texto and "pauta da" in texto.lower():
            return limpar_nome_sessao(texto)
    return "Sessão"

def extrair_pauta(caminho_arquivo):
    doc = Document(caminho_arquivo)
    lista_itens = []
    item_atual = {}
    sessao_formatada = encontrar_sessao_formatada(doc)
    regex_processo = re.compile(r"Processo\s*n[º\.]?\s*([\d\.\-\/]+)", re.IGNORECASE)

    for para in doc.paragraphs:
        texto = para.text.strip()
        match = regex_processo.search(texto)
        if match:
            if item_atual: lista_itens.append(item_atual)
            item_atual = {"processo": match.group(1), "assunto": "", "conselheiro": "", "sessao": sessao_formatada}
            continue
        if (texto.startswith("Objeto:") or texto.startswith("Assunto:")) and item_atual:
            item_atual["assunto"] = texto.split(":", 1)[1].strip()
        elif texto.startswith("Relator:") and item_atual:
            item_atual["conselheiro"] = texto.split(":", 1)[1].strip()

    if item_atual: lista_itens.append(item_atual)
    return lista_itens

def gerar_word_pauta(dados, caminho_saida):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(8)
    
    doc.add_heading('Relatório de Processos', 0)

    colunas = ["Nº Processo", "Assunto", "DISTRIBUIDO P/ CONSELHEIRO(A)", "Sessão", "Data da Assinatura", "Data da Publicação"]
    table = doc.add_table(rows=1, cols=len(colunas))
    table.style = 'Table Grid'
    table.autofit = False 

    for i, col in enumerate(colunas):
        formatar_celula_tabela(table.rows[0].cells[i], col, negrito=True)

    for item in dados:
        row = table.add_row().cells
        valores = [item['processo'], item['assunto'], item['conselheiro'], item['sessao'], "", ""]
        for i, valor in enumerate(valores):
            formatar_celula_tabela(row[i], valor, negrito=(i==0))

    doc.save(caminho_saida)

# ==============================================================================
# 3. LÓGICA DO BOTÃO 2: RESOLUÇÕES (ARIAL 12 - TEXTO CORRIDO)
# ==============================================================================

def extrair_resolucoes_dados(caminho_arquivo):
    doc = Document(caminho_arquivo)
    lista_dados = []
    ano_atual = datetime.now().year
    
    # Lê TODAS as tabelas do documento, não importa a página
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 4:
                
                raw_resolucao = row.cells[1].text.strip()
                assunto_cru = row.cells[3].text.strip()
                
                # --- FILTROS DE SEGURANÇA (Para arquivos gigantes) ---
                
                # 1. Ignora linhas onde o Assunto contém "Virtual" (cabeçalho da Pauta)
                if "virtual" in assunto_cru.lower() and len(row.cells) > 4:
                     continue 
                
                # 2. Ignora linhas onde o Assunto é muito curto (ruído/tabelas vazias)
                if len(assunto_cru) < 5:
                    continue

                # 3. Verifica se a Coluna 1 realmente parece um número
                eh_numero = any(char.isdigit() for char in raw_resolucao)
                
                # 4. Ignora se for o próprio cabeçalho da tabela ("Nº Resolução")
                if "resolução" in raw_resolucao.lower() and not eh_numero:
                    continue

                if raw_resolucao and eh_numero:
                    
                    # Limpeza e Adição do Ano
                    num_limpo = raw_resolucao.replace("Resolução", "").replace("nº", "").strip()
                    
                    if "/" in num_limpo:
                        resolucao_final = f"Resolução nº {num_limpo}"
                    else:
                        resolucao_final = f"Resolução nº {num_limpo}/{ano_atual}"

                    lista_dados.append({
                        "res": resolucao_final,
                        "assunto": assunto_cru
                    })

    return lista_dados

def gerar_word_texto_corrido(lista_dados, caminho_saida):
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    doc.add_heading('Resoluções para Publicação', 0)

    if not lista_dados:
        doc.add_paragraph("Nenhuma resolução encontrada.")
    else:
        for item in lista_dados:
            p = doc.add_paragraph()
            formatar_paragrafo_hibrido(p, item['res'], item['assunto'])

    doc.save(caminho_saida)

# ==============================================================================
# 4. INTERFACE GRÁFICA (GUI)
# ==============================================================================

def acao_pauta():
    arquivo = filedialog.askopenfilename(title="Selecione a Pauta (.docx)", filetypes=[("Word", "*.docx")])
    if not arquivo: return
    try:
        lbl_status.config(text="Processando Pauta...", fg="blue")
        root.update()
        dados = extrair_pauta(arquivo)
        if not dados:
            messagebox.showwarning("Aviso", "Nenhum processo encontrado.")
            lbl_status.config(text="Aguardando...", fg="black")
            return
        destino = filedialog.asksaveasfilename(title="Salvar Tabela Pauta", defaultextension=".docx", initialfile="Tabela_Pauta.docx")
        if not destino: return
        gerar_word_pauta(dados, destino)
        lbl_status.config(text="Pauta Concluída!", fg="green")
        messagebox.showinfo("Sucesso", f"Tabela de Pauta gerada!\n{len(dados)} processos.")
    except Exception as e:
        messagebox.showerror("Erro", str(e))
        lbl_status.config(text="Erro.", fg="red")

def acao_email():
    arquivo = filedialog.askopenfilename(title="Selecione o Doc de Resoluções (.docx)", filetypes=[("Word", "*.docx")])
    if not arquivo: return
    try:
        lbl_status.config(text="Varrendo documento inteiro...", fg="blue")
        root.update()
        
        dados = extrair_resolucoes_dados(arquivo)
        
        if not dados:
            messagebox.showwarning("Aviso", "Nenhuma resolução encontrada.\nVerifique se o arquivo contém a tabela esperada.")
            lbl_status.config(text="Aguardando...", fg="black")
            return
            
        destino = filedialog.asksaveasfilename(title="Salvar Texto Resoluções", defaultextension=".docx", initialfile="Resolucoes_Texto.docx")
        if not destino: return
        
        gerar_word_texto_corrido(dados, destino)
        lbl_status.config(text="Texto Concluído!", fg="green")
        messagebox.showinfo("Sucesso", f"Arquivo gerado!\nFormatado: Arial 12.\nEncontrados: {len(dados)} itens.")
    except Exception as e:
        messagebox.showerror("Erro", str(e))
        lbl_status.config(text="Erro.", fg="red")

# Configuração da Janela
root = tk.Tk()
root.title("Automação MP - Final v8 (Alta Capacidade)")
root.geometry("450x300")
root.configure(bg="#f0f0f0")

tk.Label(root, text="Selecione a automação desejada:", bg="#f0f0f0", font=("Arial", 12)).pack(pady=20)

btn_pauta = tk.Button(root, text="1. Pauta (Processos em Tabela)", command=acao_pauta, bg="#2b2b2b", fg="white", font=("Times New Roman", 11, "bold"), width=35, height=2)
btn_pauta.pack(pady=10)

btn_email = tk.Button(root, text="2. Resoluções (Texto Formatado)", command=acao_email, bg="#005f99", fg="white", font=("Arial", 11, "bold"), width=35, height=2)
btn_email.pack(pady=10)

lbl_status = tk.Label(root, text="Aguardando...", bg="#f0f0f0", fg="#555")
lbl_status.pack(side="bottom", pady=10)

root.mainloop()